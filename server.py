from __future__ import annotations

import argparse
import json
import os
import re
import urllib.error
import urllib.request
from html import escape
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent
HOST = "127.0.0.1"
PORT = 8765
CANONICAL_CHECKS = [
    ("mahkeme", "Mahkemeye hitap"),
    ("dava_turu", "Dava türü ve istemin belirginliği"),
    ("yd_ibaresi", "Yürütmenin durdurulması ibaresi"),
    ("davaci", "Davacı kimliği ve adresi"),
    ("vekil", "Vekil bilgisi"),
    ("davali", "Davalı idare ve husumet"),
    ("dava_konusu", "Dava konusu işlem"),
    ("kesin_yurutulebilir", "Kesin ve yürütülebilir işlem"),
    ("teblig", "Tebliğ/öğrenme tarihi"),
    ("sure", "Süre unsuru"),
    ("aciklamalar", "Açıklamalar ve hukuka aykırılık nedenleri"),
    ("hukuki_nedenler", "Hukuki nedenler"),
    ("sonuc_istem", "Sonuç ve istem"),
    ("deliller", "Deliller"),
    ("ekler", "Ekler ve dosya belgeleri"),
    ("imza_tarih", "İmza ve tarih"),
]


class PetitionHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def do_POST(self):
        if self.path == "/ai-analyze":
            self.handle_ai_analyze()
            return

        if self.path == "/export-docx":
            self.handle_export_docx()
            return

        if self.path == "/export-pdf":
            self.handle_export_pdf()
            return

        if self.path != "/extract":
            self.send_error(404, "Not found")
            return

        try:
            length = int(self.headers.get("Content-Length", "0"))
            content_type = self.headers.get("Content-Type", "")
            body = self.rfile.read(length)
            filename, data = parse_multipart_file(content_type, body)
            if not filename or not data:
                self.send_json({"error": "Dosya bulunamadı."}, status=400)
                return

            filename = Path(filename).name
            text = extract_text(filename, data)
            if not text.strip():
                self.send_json(
                    {"error": "Dosyadan okunabilir metin çıkarılamadı."},
                    status=422,
                )
                return

            self.send_json({"filename": filename, "text": text})
        except Exception as exc:
            self.send_json({"error": f"Dosya işlenemedi: {exc}"}, status=500)

    def handle_ai_analyze(self):
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            petition_text = str(payload.get("text", "")).strip()
            if not petition_text:
                self.send_json({"error": "Dilekçe metni boş."}, status=400)
                return

            analysis = analyze_with_openai(petition_text)
            self.send_json({"analysis": analysis})
        except MissingOpenAIKeyError as exc:
            self.send_json({"error": str(exc)}, status=503)
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            self.send_json({"error": f"OpenAI API hatası: {detail}"}, status=502)
        except Exception as exc:
            self.send_json({"error": f"OpenAI analizi çalıştırılamadı: {exc}"}, status=500)

    def handle_export_docx(self):
        try:
            payload = self.read_json_payload()
            body = build_docx_export(payload)
            self.send_binary(
                body,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "idari-dava-dilekce-raporu.docx",
            )
        except Exception as exc:
            self.send_json({"error": f"Word dosyası üretilemedi: {exc}"}, status=500)

    def handle_export_pdf(self):
        try:
            payload = self.read_json_payload()
            body = build_pdf_export(payload)
            self.send_binary(body, "application/pdf", "idari-dava-dilekce-raporu.pdf")
        except Exception as exc:
            self.send_json({"error": f"PDF dosyası üretilemedi: {exc}"}, status=500)

    def read_json_payload(self):
        length = int(self.headers.get("Content-Length", "0"))
        return json.loads(self.rfile.read(length).decode("utf-8"))

    def send_json(self, payload, status=200):
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_binary(self, body: bytes, content_type: str, filename: str):
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages)

    if suffix == ".docx":
        document = Document(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    if suffix == ".txt":
        return data.decode("utf-8", errors="replace")

    raise ValueError("Yalnızca PDF, DOCX ve TXT dosyaları desteklenir.")


def build_docx_export(payload: dict) -> bytes:
    analysis = payload.get("analysis") or {}
    draft = clean_export_text(payload.get("draft") or analysis.get("revisedPetition") or "")
    document = Document()
    document.add_heading("İdari Dava Dilekçesi Ön İnceleme Raporu", 0)

    add_docx_key_values(
        document,
        [
            ("Sonuç", analysis.get("verdict", "-")),
            ("Uygunluk puanı", f"{analysis.get('score', '-')}%"),
            ("Tespit edilen dava türü", analysis.get("detectedCaseType", "-")),
            ("Tespit gerekçesi", analysis.get("detectedCaseTypeReason", "-")),
        ],
    )

    document.add_heading("Kısa Değerlendirme", level=1)
    add_docx_paragraphs(document, analysis.get("summary") or "-")

    add_docx_list(document, "Eksik Bilgiler", analysis.get("missingInformation") or ["Eksik gerçek bilgi bildirilmedi."])
    add_docx_list(document, "Düzeltilebilir Noktalar", analysis.get("fixableIssues") or ["Biçimsel düzeltme önerisi bildirilmedi."])
    add_docx_list(document, "Ek/Dosya Kontrolü", analysis.get("attachmentIssues") or ["Ek/dosya kontrolü için ayrıca uyarı bildirilmedi."])

    document.add_heading("Kullanım Notu", level=1)
    add_docx_paragraphs(
        document,
        "Bu rapor ön inceleme desteği sağlar; nihai hukuki değerlendirme ve dava stratejisi somut dosya üzerinden ayrıca incelenmelidir.",
    )

    document.add_heading("Detaylı Kontrol Tablosu", level=1)
    checklist = analysis.get("checklist") or []
    if checklist:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["Durum", "Unsur", "Dayanak", "Öneri"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in checklist:
            cells = table.add_row().cells
            cells[0].text = clean_export_text(item.get("status") or "-")
            cells[1].text = clean_export_text(item.get("title") or "-")
            cells[2].text = clean_export_text(item.get("evidence") or "-")
            cells[3].text = clean_export_text(item.get("recommendation") or item.get("explanation") or "-")
    else:
        document.add_paragraph("Kontrol tablosu oluşturulamadı.")

    document.add_heading("Düzeltilmiş Taslak", level=1)
    add_docx_paragraphs(document, draft or "[Taslak üretilemedi.]")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def add_docx_key_values(document: Document, values: list[tuple[str, object]]) -> None:
    for key, value in values:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{key}: ").bold = True
        paragraph.add_run(clean_export_text(value))


def add_docx_list(document: Document, title: str, values: list[str]) -> None:
    document.add_heading(title, level=1)
    for value in values:
        document.add_paragraph(clean_export_text(value), style="List Bullet")


def add_docx_paragraphs(document: Document, text: object) -> None:
    parts = split_export_paragraphs(clean_export_text(text))
    for part in parts or ["-"]:
        document.add_paragraph(part)


def build_pdf_export(payload: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = register_pdf_font(pdfmetrics, TTFont)
    styles = getSampleStyleSheet()
    for style_name in ["Title", "Heading1", "Heading2", "Normal", "BodyText"]:
        styles[style_name].fontName = font_name
    styles["Normal"].fontSize = 9
    styles["Normal"].leading = 12

    analysis = payload.get("analysis") or {}
    draft = clean_export_text(payload.get("draft") or analysis.get("revisedPetition") or "")
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=1.4 * cm,
        leftMargin=1.4 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )
    story = []

    story.append(Paragraph("İdari Dava Dilekçesi Ön İnceleme Raporu", styles["Title"]))
    story.append(Spacer(1, 10))
    for key, value in [
        ("Sonuç", analysis.get("verdict", "-")),
        ("Uygunluk puanı", f"{analysis.get('score', '-')}%"),
        ("Tespit edilen dava türü", analysis.get("detectedCaseType", "-")),
        ("Tespit gerekçesi", analysis.get("detectedCaseTypeReason", "-")),
    ]:
        story.append(Paragraph(f"<b>{escape(key)}:</b> {escape(clean_export_text(value))}", styles["Normal"]))
    story.append(Spacer(1, 10))

    add_pdf_section(story, styles, "Kısa Değerlendirme", analysis.get("summary") or "-")
    add_pdf_list(story, styles, "Eksik Bilgiler", analysis.get("missingInformation") or ["Eksik gerçek bilgi bildirilmedi."])
    add_pdf_list(story, styles, "Düzeltilebilir Noktalar", analysis.get("fixableIssues") or ["Biçimsel düzeltme önerisi bildirilmedi."])
    add_pdf_list(story, styles, "Ek/Dosya Kontrolü", analysis.get("attachmentIssues") or ["Ek/dosya kontrolü için ayrıca uyarı bildirilmedi."])
    add_pdf_section(
        story,
        styles,
        "Kullanım Notu",
        "Bu rapor ön inceleme desteği sağlar; nihai hukuki değerlendirme ve dava stratejisi somut dosya üzerinden ayrıca incelenmelidir.",
    )

    checklist = analysis.get("checklist") or []
    story.append(Paragraph("Detaylı Kontrol Tablosu", styles["Heading1"]))
    if checklist:
        data = [["Durum", "Unsur", "Dayanak", "Öneri"]]
        for item in checklist:
            data.append([
                clean_export_text(item.get("status") or "-"),
                clean_export_text(item.get("title") or "-"),
                clean_export_text(item.get("evidence") or "-"),
                clean_export_text(item.get("recommendation") or item.get("explanation") or "-"),
            ])
        table = Table(data, colWidths=[2.0 * cm, 3.6 * cm, 5.0 * cm, 5.2 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF8")),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B7C3D6")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)
    else:
        story.append(Paragraph("Kontrol tablosu oluşturulamadı.", styles["Normal"]))

    story.append(Spacer(1, 12))
    add_pdf_section(story, styles, "Düzeltilmiş Taslak", draft or "[Taslak üretilemedi.]")

    document.build(story)
    return output.getvalue()


def register_pdf_font(pdfmetrics, TTFont) -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            pdfmetrics.registerFont(TTFont("AppFont", candidate))
            return "AppFont"
    return "Helvetica"


def add_pdf_section(story: list, styles, title: str, text: object) -> None:
    from reportlab.platypus import Paragraph, Spacer

    story.append(Paragraph(escape(title), styles["Heading1"]))
    for paragraph in split_export_paragraphs(clean_export_text(text)) or ["-"]:
        story.append(Paragraph(escape(paragraph), styles["Normal"]))
        story.append(Spacer(1, 4))


def add_pdf_list(story: list, styles, title: str, values: list[str]) -> None:
    from reportlab.platypus import Paragraph, Spacer

    story.append(Paragraph(escape(title), styles["Heading1"]))
    for value in values:
        story.append(Paragraph(f"• {escape(clean_export_text(value))}", styles["Normal"]))
    story.append(Spacer(1, 8))


def clean_export_text(value: object) -> str:
    return re.sub(r"[ \t]+", " ", str(value or "")).strip()


def split_export_paragraphs(text: str) -> list[str]:
    text = text.replace("\r", "\n")
    return [part.strip() for part in re.split(r"\n{2,}|\n", text) if part.strip()]


class MissingOpenAIKeyError(RuntimeError):
    pass


def analyze_with_openai(petition_text: str) -> dict:
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise MissingOpenAIKeyError(
            "OPENAI_API_KEY tanımlı değil. Render Environment bölümüne OpenAI API anahtarı eklenmeli."
        )

    model = os.environ.get("OPENAI_MODEL", "gpt-5.4-mini")
    evidence = extract_local_evidence(petition_text)
    prompt = build_legal_prompt(petition_text, evidence)
    body = {
        "model": model,
        "input": prompt,
        "text": {
            "format": {
                "type": "json_schema",
                "name": "iyuk_petition_analysis",
                "strict": True,
                "schema": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "verdict": {"type": "string", "enum": ["Geçer", "Riskli", "Geçmez"]},
                        "score": {"type": "integer", "minimum": 0, "maximum": 100},
                        "detectedCaseType": {"type": "string"},
                        "detectedCaseTypeReason": {"type": "string"},
                        "summary": {"type": "string"},
                        "checklist": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "id": {"type": "string"},
                                    "title": {"type": "string"},
                                    "status": {"type": "string", "enum": ["uygun", "riskli", "eksik", "düzeltilmeli"]},
                                    "evidence": {"type": "string"},
                                    "explanation": {"type": "string"},
                                    "recommendation": {"type": "string"},
                                },
                                "required": ["id", "title", "status", "evidence", "explanation", "recommendation"],
                            },
                        },
                        "missingInformation": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "fixableIssues": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "attachmentIssues": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                        "revisedPetition": {"type": "string"},
                    },
                    "required": [
                        "verdict",
                        "score",
                        "detectedCaseType",
                        "detectedCaseTypeReason",
                        "summary",
                        "checklist",
                        "missingInformation",
                        "fixableIssues",
                        "attachmentIssues",
                        "revisedPetition",
                    ],
                },
            }
        },
    }

    request = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(body).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=90) as response:
        data = json.loads(response.read().decode("utf-8"))

    output_text = extract_response_text(data)
    analysis = json.loads(output_text)
    analysis = normalize_checklist(analysis, evidence)
    analysis = enforce_mechanical_findings(analysis, evidence)
    return normalize_issue_categories(analysis, evidence)


def build_legal_prompt(petition_text: str, evidence: dict) -> str:
    return f"""
Sen idari yargılama usulü alanında çalışan, İYUK m.3 ve idari dava dilekçelerinin ön inceleme şartları bakımından uzmanlaşmış bir dilekçe kontrol motorusun.

Görev:
- Kullanıcı dava türünü bilmek zorunda değildir. Önce dilekçenin türünü içerikten kendin belirle: iptal davası, tam yargı davası, iptal + tam yargı davası, yürütmenin durdurulması talepli dava veya belirsiz/karma.
- Tespit ettiğin dava türünü ve gerekçesini yaz.
- İYUK m.3 kapsamındaki unsurları tek tek değerlendir.
- Görev, yetki, süre, ehliyet, husumet, dava konusu işlem, kesin/yürütülebilir işlem, tebliğ/öğrenme tarihi, sonuç ve istem, deliller, ekler, imza/tarih ve dava türüne özgü biçimsel unsurlar bakımından riskleri yaz.
- Dilekçede yürütmenin durdurulması isteniyor veya olayın niteliği YD talebini gerektiriyor gibi görünüyorsa, “YÜRÜTMENİN DURDURULMASI TALEPLİDİR” ibaresinin bulunup bulunmadığını özellikle kontrol et.
- Kullanıcı iptal davası gibi görünen bir dilekçe sunmuş ama YD talebine ilişkin olgular/istemler var ve ibare eksikse bunu açıkça eksik/riskli unsur olarak işaretle.
- Eksik gerçek bilgileri uydurma. Eksik bilgi gereken yerlere köşeli parantezli açıklama koy.
- missingInformation alanına sadece dilekçe metninde bulunmayan ve kullanıcıdan öğrenilmesi gereken gerçek olay/bilgi eksiklerini yaz.
- missingInformation alanına dava açılış tarihi yazma; dava açılış tarihi mahkemeye sunumla oluşur. Dilekçe sonunda tarih yoksa bunu checklist içinde “imza ve tarih” başlığında düzeltilmeli/riskli değerlendir.
- Deliller veya belgeler dilekçede sayılmışsa ama ayrı “EKLER” başlığı yoksa bunu kritik eksik yapma; “düzeltilmeli” olarak sınıflandır ve fixableIssues/attachmentIssues alanına koy.
- Tebliğ belgesi, vekaletname veya eklerin fiilen dosyaya konulup konulmadığı metinden anlaşılamıyorsa bunu missingInformation değil attachmentIssues olarak yaz.
- Husumet/yetki bakımından açık çelişki yoksa “riskli” deme; “uygun” veya gerekirse “düzeltilmeli” deyip dosya ekiyle teyit öner.
- Kritik eksik yalnızca mahkeme, davacı, davalı, dava konusu işlem, tebliğ/öğrenme tarihi, sonuç/istem gibi ön incelemede doğrudan sorun yaratabilecek zorunlu unsurlar gerçekten yoksa kullanılmalıdır.
- Düzeltilebilen anlatım, başlık, konu, sonuç ve istem bölümlerini uygun dilekçe formuna getir.
- Yürütmenin durdurulması talepli olduğu sonucuna varırsan ilgili ibareyi büyük harfli ve belirgin şekilde taslağa ekle.
- Her kontrol maddesinde kararını dilekçedeki somut metin parçasına bağla.
- Dayanak alanına mümkün olduğunca dilekçeden aynen kısa alıntı yaz. Yorum, özet veya uydurma metin yazma.
- Bir unsur için “uygun” diyebilmen için dilekçede o unsuru açıkça gösteren bir metin parçası bulunmalıdır.
- Metinde açık karşılığı yoksa “eksik” veya “riskli” de; tahminle uygun deme.
- Aynı dilekçe aynı kanıtlarla değerlendirildiğinde aynı sonuca varacak şekilde tutarlı davran.
- Kendi ilk kararını ayrıca denetle: “uygun” dediğin her maddenin dayanağı gerçekten dilekçede var mı, yoksa durumu riskli/eksik yap.
- Aşağıdaki yerel kanıt çıkarımını yardımcı veri olarak kullan; ancak nihai kararı dilekçenin tamamına göre ver.
- Cevabı yalnızca istenen JSON şemasına uygun ver.

Yerel kanıt çıkarımı:
{json.dumps(evidence, ensure_ascii=False, indent=2)}

Dilekçe metni:
{petition_text}
""".strip()


def extract_response_text(data: dict) -> str:
    if data.get("output_text"):
        return data["output_text"]

    chunks = []
    for output in data.get("output", []):
        for content in output.get("content", []):
            text = content.get("text")
            if text:
                chunks.append(text)
    if chunks:
        return "\n".join(chunks)

    raise ValueError("OpenAI yanıtında metin bulunamadı.")


def enforce_mechanical_findings(analysis: dict, evidence: dict) -> dict:
    checklist = analysis.get("checklist", [])
    evidence_requirements = {
        "court": ("mahkemeye_hitap", "Mahkemeye hitap bölümü mekanik kontrolde tespit edilemedi."),
        "mahkemeye_hitap": ("mahkemeye_hitap", "Mahkemeye hitap bölümü mekanik kontrolde tespit edilemedi."),
        "plaintiff": ("davacı", "Davacı bölümü mekanik kontrolde tespit edilemedi."),
        "davacı": ("davacı", "Davacı bölümü mekanik kontrolde tespit edilemedi."),
        "defendant": ("davali", "Davalı idare bölümü mekanik kontrolde tespit edilemedi."),
        "davalı": ("davali", "Davalı idare bölümü mekanik kontrolde tespit edilemedi."),
        "subject": ("konu", "Konu bölümü mekanik kontrolde tespit edilemedi."),
        "konu": ("konu", "Konu bölümü mekanik kontrolde tespit edilemedi."),
        "notice": ("teblig_ogrenme_tarihi", "Tebliğ/öğrenme tarihi mekanik kontrolde tespit edilemedi."),
        "teblig": ("teblig_ogrenme_tarihi", "Tebliğ/öğrenme tarihi mekanik kontrolde tespit edilemedi."),
        "tebliğ": ("teblig_ogrenme_tarihi", "Tebliğ/öğrenme tarihi mekanik kontrolde tespit edilemedi."),
        "amount": ("tazminat_miktari", "Tam yargı bakımından tazminat miktarı mekanik kontrolde tespit edilemedi."),
        "miktar": ("tazminat_miktari", "Tam yargı bakımından tazminat miktarı mekanik kontrolde tespit edilemedi."),
        "request": ("sonuc_istem", "Sonuç ve istem bölümü mekanik kontrolde tespit edilemedi."),
        "sonuç": ("sonuc_istem", "Sonuç ve istem bölümü mekanik kontrolde tespit edilemedi."),
        "istem": ("sonuc_istem", "Sonuç ve istem bölümü mekanik kontrolde tespit edilemedi."),
        "annex": ("ekler_deliller", "Ekler/deliller bölümü mekanik kontrolde tespit edilemedi."),
        "ek": ("ekler_deliller", "Ekler/deliller bölümü mekanik kontrolde tespit edilemedi."),
        "delil": ("ekler_deliller", "Ekler/deliller bölümü mekanik kontrolde tespit edilemedi."),
        "signature": ("imza_tarih", "İmza/tarih alanı mekanik kontrolde tespit edilemedi."),
        "imza": ("imza_tarih", "İmza/tarih alanı mekanik kontrolde tespit edilemedi."),
    }

    for item in checklist:
        status = fold_tr(str(item.get("status", "")))
        item_evidence = str(item.get("evidence", "")).strip()
        if status != "uygun":
            continue

        if not item_evidence or fold_tr(item_evidence) in {"-", "yok", "bulunamadı"}:
            downgrade_item(item, "Uygunluk dayanağı gösterilmediği için bu unsur riskli kabul edildi.")
            continue

        requirement = find_requirement(item, evidence_requirements)
        if requirement:
            evidence_key, note = requirement
            if not evidence_value_present(evidence_key, evidence):
                downgrade_item(item, note)

    missing_count = sum(1 for item in checklist if fold_tr(str(item.get("status", ""))) != "uygun")
    if checklist:
        recalculated = round(((len(checklist) - missing_count) / len(checklist)) * 100)
        current_score = int(analysis.get("score", recalculated))
        analysis["score"] = current_score if 0 <= current_score <= 100 else recalculated

    if missing_count >= max(4, len(checklist) // 3):
        analysis["verdict"] = "Geçmez"
    elif missing_count > 0 and analysis.get("verdict") == "Geçer":
        analysis["verdict"] = "Riskli"

    return analysis


def normalize_checklist(analysis: dict, evidence: dict) -> dict:
    source_items = analysis.get("checklist", [])
    normalized_items = []
    for check_id, title in CANONICAL_CHECKS:
        model_item = find_model_item(check_id, title, source_items)
        mechanical = mechanical_item(check_id, title, evidence)
        normalized_items.append(merge_item(model_item, mechanical))

    analysis["checklist"] = normalized_items
    missing_count = sum(1 for item in normalized_items if fold_tr(item["status"]) in {"eksik", "riskli"})
    fixable_count = sum(1 for item in normalized_items if fold_tr(item["status"]) == "düzeltilmeli")
    analysis["score"] = max(0, min(100, round(((len(normalized_items) - missing_count - (fixable_count * 0.35)) / len(normalized_items)) * 100)))
    if missing_count == 0:
        analysis["verdict"] = "Geçer" if fixable_count <= 3 else "Riskli"
    elif missing_count <= 2:
        analysis["verdict"] = "Riskli"
    else:
        analysis["verdict"] = "Geçmez"
    return analysis


def find_model_item(check_id: str, title: str, items: list[dict]) -> dict | None:
    haystacks = {
        "mahkeme": ["mahkeme", "hitap", "görev", "yetki"],
        "dava_turu": ["dava tür", "istem", "iptal", "tam yargı"],
        "yd_ibaresi": ["yürütme", "yd"],
        "davaci": ["davacı", "kimlik", "adres", "ehliyet"],
        "vekil": ["vekil", "avukat", "vekâlet", "vekalet"],
        "davali": ["davalı", "husumet"],
        "dava_konusu": ["dava konusu", "işlem"],
        "kesin_yurutulebilir": ["kesin", "yürütülebilir", "icrai"],
        "teblig": ["tebliğ", "tebellüğ", "öğrenme"],
        "sure": ["süre"],
        "aciklamalar": ["açıklama", "vakıa", "sebep"],
        "hukuki_nedenler": ["hukuki", "neden", "sebep"],
        "sonuc_istem": ["sonuç", "istem", "talep"],
        "deliller": ["delil"],
        "ekler": ["ekler", "ek "],
        "imza_tarih": ["imza", "tarih"],
    }
    keywords = haystacks.get(check_id, [title])
    for item in items:
        text = fold_tr(f"{item.get('id', '')} {item.get('title', '')}")
        if any(fold_tr(keyword) in text for keyword in keywords):
            return item
    return None


def merge_item(model_item: dict | None, mechanical: dict) -> dict:
    if not model_item:
        return mechanical

    model_status = fold_tr(str(model_item.get("status", "")))
    mechanical_status = fold_tr(mechanical["status"])
    final = mechanical.copy()

    if mechanical_status in {"eksik", "riskli", "düzeltilmeli"}:
        return final

    if mechanical_status == "uygun" and model_status in {"riskli", "eksik", "düzeltilmeli"}:
        final["status"] = "düzeltilmeli" if model_status == "düzeltilmeli" else "riskli"
        final["explanation"] = model_item.get("explanation") or final["explanation"]
        final["recommendation"] = model_item.get("recommendation") or final["recommendation"]
    else:
        final["explanation"] = model_item.get("explanation") or final["explanation"]
        final["recommendation"] = model_item.get("recommendation") or final["recommendation"]

    final["evidence"] = mechanical.get("evidence") or model_item.get("evidence") or "-"
    return final


def mechanical_item(check_id: str, title: str, evidence: dict) -> dict:
    item = {
        "id": check_id,
        "title": title,
        "status": "uygun",
        "evidence": "-",
        "explanation": "Mekanik kontrolde uygunluk dayanağı tespit edildi.",
        "recommendation": "Korunabilir.",
    }

    def set_item(status: str, key: str, explanation: str, recommendation: str) -> dict:
        item["status"] = status
        item["evidence"] = str(evidence.get(key, "") or "-")
        item["explanation"] = explanation
        item["recommendation"] = recommendation
        return item

    if check_id == "mahkeme":
        return set_item("uygun" if evidence.get("mahkemeye_hitap") else "eksik", "mahkemeye_hitap", "Mahkemeye hitap kontrol edildi.", "Mahkeme başlığı açıkça yazılmalıdır.")
    if check_id == "dava_turu":
        return set_item("uygun" if evidence.get("konu") or evidence.get("sonuc_istem") else "eksik", "konu", "Dava türü konu/istem bölümünden anlaşılabilir.", "Dava türü konu ve istemde netleştirilmelidir.")
    if check_id == "yd_ibaresi":
        needs_yd = bool(evidence.get("yd_ihtiyaci") or evidence.get("yd_ibaresi"))
        status = "uygun" if evidence.get("yd_ibaresi") else "eksik" if needs_yd else "uygun"
        return set_item(status, "yd_ibaresi", "YD ibaresi ve YD ihtiyacı kontrol edildi.", "YD talebi varsa ibare büyük harfle yazılmalıdır.")
    if check_id == "davaci":
        ok = evidence_value_present("davacı_kimlik_adres", evidence)
        return set_item("uygun" if ok else "eksik", "davacı", "Davacı kimliği ve adresi kontrol edildi.", "Davacı adı, TCKN ve açık adres yazılmalıdır.")
    if check_id == "vekil":
        status = "uygun" if evidence.get("vekil") else "uygun"
        return set_item(status, "vekil", "Vekil bilgisi varsa kontrol edildi.", "Vekil varsa vekaletname ek/dosya kontrolünde teyit edilmelidir.")
    if check_id == "davali":
        return set_item("uygun" if evidence.get("davali") else "eksik", "davali", "Davalı idare kontrol edildi.", "Davalı idare açıkça gösterilmelidir.")
    if check_id == "dava_konusu":
        return set_item("uygun" if evidence.get("konu") else "eksik", "konu", "Dava konusu işlem kontrol edildi.", "İşlem tarihi/sayısı ve istem açıkça yazılmalıdır.")
    if check_id == "kesin_yurutulebilir":
        return set_item("uygun" if evidence.get("islem_gorunumu") else "riskli", "islem_gorunumu", "İşlemin icrai görünümü kontrol edildi.", "Dava konusu işlemin kesin ve yürütülebilir olduğu metinden anlaşılmalıdır.")
    if check_id == "teblig":
        return set_item("uygun" if evidence.get("teblig_ogrenme_tarihi") else "eksik", "teblig_ogrenme_tarihi", "Tebliğ/öğrenme tarihi kontrol edildi.", "Tebliğ veya öğrenme tarihi yazılmalıdır.")
    if check_id == "sure":
        return set_item("uygun" if evidence.get("teblig_ogrenme_tarihi") else "riskli", "teblig_ogrenme_tarihi", "Süre için başlangıç tarihi kontrol edildi.", "Süre hesabı tebliğ/öğrenme tarihine göre ayrıca yapılmalıdır.")
    if check_id == "aciklamalar":
        return set_item("uygun" if evidence.get("aciklamalar") else "riskli", "aciklamalar", "Maddi olaylar kontrol edildi.", "Maddi olaylar ve hukuka aykırılık nedenleri açıklanmalıdır.")
    if check_id == "hukuki_nedenler":
        return set_item("uygun" if evidence.get("hukuki_nedenler") else "düzeltilmeli", "hukuki_nedenler", "Hukuki nedenler kontrol edildi.", "Hukuki nedenler başlığı eklenebilir veya güçlendirilebilir.")
    if check_id == "sonuc_istem":
        return set_item("uygun" if evidence.get("sonuc_istem") else "eksik", "sonuc_istem", "Sonuç ve istem kontrol edildi.", "Sonuç ve istem açıkça yazılmalıdır.")
    if check_id == "deliller":
        return set_item("uygun" if evidence.get("deliller") or evidence.get("ekler_deliller") else "düzeltilmeli", "deliller", "Deliller kontrol edildi.", "Deliller ayrı başlık altında gösterilmelidir.")
    if check_id == "ekler":
        if evidence.get("ekler"):
            return set_item("uygun", "ekler", "Ekler başlığı mevcut.", "Ekler numaralandırılabilir.")
        return set_item("düzeltilmeli" if evidence.get("deliller") else "eksik", "deliller", "Deliller/ekler sunuluş biçimi kontrol edildi.", "Ayrı EKLER başlığı açılıp belgeler numaralandırılmalıdır.")
    if check_id == "imza_tarih":
        if evidence.get("imza_tarih"):
            return set_item("uygun", "imza_tarih", "İmza ve tarih mevcut.", "Korunabilir.")
        if evidence.get("imza_var"):
            return set_item("düzeltilmeli", "imza_var", "İmza var ancak dilekçe tarihi tespit edilemedi.", "Dilekçe sonuna gün/ay/yıl tarihi eklenmelidir.")
        return set_item("eksik", "imza_var", "İmza/tarih alanı tespit edilemedi.", "Dilekçe tarih ve imza ile tamamlanmalıdır.")

    return item


def normalize_issue_categories(analysis: dict, evidence: dict) -> dict:
    missing = []
    fixable = []
    attachments = []

    for item in analysis.get("checklist", []):
        title = fold_tr(str(item.get("title", "")))
        status = fold_tr(str(item.get("status", "")))
        recommendation = str(item.get("recommendation", "")).strip()

        if ("ek" in title or "delil" in title) and status == "eksik" and evidence.get("ekler_deliller"):
            item["status"] = "düzeltilmeli"
            item["explanation"] = (
                f"{item.get('explanation', '')} Delil/ek bilgisi metinde bulunduğu için bu husus kritik eksik değil, biçimsel tamamlama olarak değerlendirildi."
            ).strip()
            if recommendation:
                attachments.append(recommendation)

        if ("husumet" in title or "yetki" in title) and status == "riskli":
            if evidence.get("davali") and evidence.get("mahkemeye_hitap"):
                item["status"] = "düzeltilmeli"
                item["explanation"] = (
                    f"{item.get('explanation', '')} Açık çelişki tespit edilmediği için bu husus kritik risk değil, dosya ekiyle teyit edilmesi gereken nokta olarak sınıflandırıldı."
                ).strip()
                if recommendation:
                    fixable.append(recommendation)

        if status in {"eksik", "riskli"} and item.get("id") in {
            "mahkeme",
            "davaci",
            "davali",
            "dava_konusu",
            "teblig",
            "sonuc_istem",
            "yd_ibaresi",
        }:
            missing.append(f"{item.get('title')}: {item.get('recommendation')}")

        if status == "düzeltilmeli":
            fixable.append(f"{item.get('title')}: {item.get('recommendation')}")

    if evidence.get("vekil"):
        attachments.append("Vekaletname dosyada bulunmalı.")
    if evidence.get("teblig_ogrenme_tarihi"):
        attachments.append("Tebliğ/öğrenme belgesi dosyada bulunmalı.")
    if evidence.get("konu") or evidence.get("islem_gorunumu"):
        attachments.append("Dava konusu işlem örneği dosyada bulunmalı.")
    if evidence.get("deliller") and not evidence.get("ekler"):
        fixable.append("Ekler ve dosya belgeleri: Ayrı EKLER başlığı açılıp belgeler numaralandırılmalıdır.")

    analysis["missingInformation"] = unique_keep_order(missing)
    analysis["fixableIssues"] = unique_keep_order(fixable)
    analysis["attachmentIssues"] = unique_keep_order(attachments)

    missing_count = sum(1 for item in analysis.get("checklist", []) if fold_tr(str(item.get("status", ""))) in {"eksik", "riskli"})
    if missing_count == 0 and analysis.get("verdict") != "Geçer":
        analysis["verdict"] = "Geçer"
    elif missing_count <= 2 and analysis.get("verdict") == "Geçmez":
        analysis["verdict"] = "Riskli"

    return analysis


def unique_keep_order(values: list[str]) -> list[str]:
    seen = set()
    result = []
    for value in values:
        cleaned = value.strip()
        if not cleaned:
            continue
        key = fold_tr(cleaned)
        if key in seen:
            continue
        seen.add(key)
        result.append(cleaned)
    return result


def find_requirement(item: dict, requirements: dict[str, tuple[str, str]]) -> tuple[str, str] | None:
    text = fold_tr(f"{item.get('id', '')} {item.get('title', '')}")
    if "davacı" in text and ("adres" in text or "kimlik" in text or "tc" in text or "t.c" in text):
        return ("davacı_kimlik_adres", "Davacı kimlik/adres bilgisi mekanik kontrolde birlikte tespit edilemedi.")
    if "davacı" in text and "adres" in text:
        return ("davacı_adresi_olabilir", "Davacı adresi mekanik kontrolde tespit edilemedi.")
    for keyword, requirement in requirements.items():
        if keyword in text:
            return requirement
    return None


def evidence_value_present(key: str, evidence: dict) -> bool:
    if key == "davacı_kimlik_adres":
        return bool(evidence.get("davacı")) and bool(evidence.get("davacı_kimlik_no")) and bool(evidence.get("davacı_adresi_olabilir"))
    value = str(evidence.get(key, "")).strip()
    return bool(value)


def downgrade_item(item: dict, note: str) -> None:
    item["status"] = "riskli"
    explanation = str(item.get("explanation", "")).strip()
    item["explanation"] = f"{explanation} {note}".strip()


def extract_local_evidence(petition_text: str) -> dict:
    sections = extract_sections(petition_text)
    text = re.sub(r"\s+", " ", petition_text).strip()
    plaintiff_section = sections.get("davacı") or find_labeled_excerpt(text, "davacı", ["vekili", "davalı", "konu"])
    signature_block = extract_signature_block(petition_text)
    return {
        "mahkemeye_hitap": sections.get("mahkeme") or find_excerpt(text, r"(danıştay|idare mahkemesi|vergi mahkemesi).{0,110}(başkanlığı'?na|dairesi'?ne)"),
        "davacı": plaintiff_section,
        "davacı_kimlik_no": find_excerpt(plaintiff_section, r"(t\.?\s*c\.?\s*)?(kimlik\s*)?(no|numarası)?\s*:?\s*[1-9][0-9]{10}"),
        "davacı_adresi_olabilir": find_excerpt(plaintiff_section, r"(adres\s*:|mahallesi|mah\.|sokak|sok\.|cadde|cad\.|bulvar|bulv\.|no\s*:|daire|d\s*:|/\s*[A-ZÇĞİÖŞÜa-zçğıöşü]+)"),
        "vekil": sections.get("vekili") or find_labeled_excerpt(text, "vekili", ["davalı", "konu", "tebellüğ", "tebliğ", "açıklamalar"]),
        "davali": sections.get("davalı") or find_labeled_excerpt(text, "davalı", ["konu", "tebellüğ", "tebliğ", "açıklamalar"]),
        "konu": sections.get("konu") or find_labeled_excerpt(text, "konu", ["tebellüğ", "tebliğ", "açıklamalar", "olaylar"]),
        "teblig_ogrenme_tarihi": sections.get("tebliğ") or find_excerpt(text, r"(tebellüğ|tebliğ|öğrenme|bildirim).{0,80}\d{1,2}[./]\d{1,2}[./]\d{4}"),
        "tazminat_miktari": find_excerpt(text, r"\d[\d.,]*\s*(tl|₺|türk lirası)"),
        "yd_ibaresi": find_excerpt(text, r"YÜRÜTMENİN\s+DURDURULMASI\s+TALEPLİDİR|yürütmenin\s+durdurulması"),
        "yd_ihtiyaci": find_excerpt(text, r"telafisi\s+güç|yürütmenin\s+durdurulması|açıkça\s+hukuka\s+aykırı"),
        "islem_gorunumu": find_excerpt(text, r"(işlem|karar|atama|naklen|ret|tesis).{0,100}(tarih|sayılı|iptal|tebliğ|atan)"),
        "aciklamalar": sections.get("açıklamalar") or find_excerpt(text, r"(açıklamalar|olaylar).{0,500}"),
        "hukuki_nedenler": sections.get("hukuki sebepler") or find_excerpt(text, r"(hukuki\s+(nedenler|sebepler)|2577|iyuk|anayasa|kanunu)"),
        "deliller": sections.get("deliller") or find_excerpt(text, r"(deliller|tanık|belge|kararı|cetveli|kayıt örneği)"),
        "ekler": sections.get("ekler"),
        "sonuc_istem": sections.get("sonuç") or find_excerpt(text, r"(sonuç|netice).{0,30}(talep|istem)"),
        "ekler_deliller": sections.get("ekler") or find_excerpt(text, r"(ekler|deliller|ek\s*:)"),
        "imza_var": find_excerpt(signature_block, r"(imza|davacı|vekili|av\.)"),
        "imza_tarih": find_excerpt(signature_block, r"\d{1,2}[./]\d{1,2}[./]\d{4}|\[?gün/ay/yıl\]?|\[?tarih\]?"),
    }


def find_excerpt(text: str, pattern: str) -> str:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return ""
    start = max(0, match.start() - 60)
    end = min(len(text), match.end() + 100)
    return text[start:end].strip()


def find_labeled_excerpt(text: str, label: str, stop_labels: list[str] | None = None) -> str:
    match = re.search(
        rf"(?:^|\s){label}\s*:?(?=\s|$)",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        return ""
    start = match.start()
    stops = stop_labels or ["vekili", "davalı", "konu", "tebellüğ", "tebliğ", "açıklamalar", "hukuki sebepler", "sonuç"]
    stop_pattern = "|".join(re.escape(stop) for stop in stops)
    next_label = re.search(
        rf"\b({stop_pattern})\s*:?(?=\s|$)",
        text[match.end() :],
        flags=re.IGNORECASE,
    )
    end = match.end() + next_label.start() if next_label else min(len(text), start + 450)
    return text[start:end].strip()


def extract_signature_block(petition_text: str) -> str:
    lines = [line.strip() for line in petition_text.replace("\r", "\n").split("\n") if line.strip()]
    start_index = max(0, len(lines) - 8)
    for index in range(len(lines) - 1, -1, -1):
        if re.search(r"^(davacı|vekili|av\.|imza)\b", lines[index], flags=re.IGNORECASE):
            start_index = max(0, index - 2)
            break
    return " ".join(lines[start_index:])


def extract_sections(petition_text: str) -> dict[str, str]:
    lines = [line.strip() for line in petition_text.replace("\r", "\n").split("\n")]
    lines = [line for line in lines if line]
    sections: dict[str, list[str]] = {}
    current_key = "mahkeme"
    sections[current_key] = []

    for line in lines:
        heading = detect_heading(line)
        if heading:
            current_key = heading
            sections.setdefault(current_key, [])
            remainder = strip_heading(line, heading)
            if remainder:
                sections[current_key].append(remainder)
            continue
        sections.setdefault(current_key, []).append(line)

    return {key: " ".join(value).strip() for key, value in sections.items() if " ".join(value).strip()}


def detect_heading(line: str) -> str | None:
    cleaned = normalize_heading(line)
    aliases = {
        "davacı": ["davacı", "davacı bilgileri", "davacı taraf", "davacılar"],
        "vekili": ["vekili", "davacı vekili", "vekil", "vekiller"],
        "davalı": ["davalı", "davalı idare", "davalı taraf", "davalılar"],
        "konu": ["konu", "dava konusu", "davanın konusu"],
        "açıklamalar": ["açıklamalar", "olaylar", "izahlar", "maddi olaylar"],
        "hukuki sebepler": ["hukuki sebepler", "hukuki nedenler", "yasal sebepler"],
        "tebliğ": ["tebellüğ tarihi", "tebliğ tarihi", "öğrenme tarihi", "bildirim tarihi"],
        "sonuç": ["sonuç ve talep", "sonuç ve istem", "netice ve talep", "sonuç", "istem"],
        "deliller": ["deliller", "deliller ve belgeler", "delil listesi"],
        "ekler": ["ekler", "ek", "ekler ve deliller"],
    }
    for key, names in aliases.items():
        for name in names:
            if cleaned == name or re.match(rf"^{re.escape(name)}\s*:", cleaned):
                return key
    return None


def strip_heading(line: str, heading: str) -> str:
    aliases = {
        "davacı": r"davacı(?:lar| bilgileri| taraf)?",
        "vekili": r"(?:davacı\s+)?vekil(?:i|ler)?",
        "davalı": r"davalı(?:lar| idare| taraf)?",
        "konu": r"(?:dava(?:nın)?\s+)?konu(?:su)?",
        "açıklamalar": r"açıklamalar|olaylar|izahlar|maddi olaylar",
        "hukuki sebepler": r"hukuki sebepler|hukuki nedenler|yasal sebepler",
        "tebliğ": r"tebellüğ tarihi|tebliğ tarihi|öğrenme tarihi|bildirim tarihi",
        "sonuç": r"sonuç\s+ve\s+(?:talep|istem)|netice\s+ve\s+talep|sonuç|istem",
        "deliller": r"deliller\s+ve\s+belgeler|delil listesi|deliller",
        "ekler": r"ekler\s+ve\s+deliller|ekler|ek",
    }
    pattern = aliases.get(heading)
    if not pattern:
        return ""
    return re.sub(rf"^\s*{pattern}\s*:?\s*", "", line, flags=re.IGNORECASE).strip()


def normalize_heading(value: str) -> str:
    value = re.sub(r"[:：]+$", "", value.strip())
    value = re.sub(r"\s+", " ", value)
    return fold_tr(value)


def fold_tr(value: str) -> str:
    return value.translate(str.maketrans({"I": "ı", "İ": "i"})).casefold()


def parse_multipart_file(content_type: str, body: bytes) -> tuple[str, bytes]:
    boundary_match = re.search(r"boundary=(?P<boundary>[^;]+)", content_type)
    if not boundary_match:
        raise ValueError("Geçersiz dosya yükleme isteği.")

    boundary = boundary_match.group("boundary").strip('"')
    delimiter = f"--{boundary}".encode("utf-8")
    for part in body.split(delimiter):
        part = part.strip()
        if not part or part == b"--":
            continue

        header_bytes, separator, file_bytes = part.partition(b"\r\n\r\n")
        if not separator:
            continue

        headers = header_bytes.decode("utf-8", errors="replace")
        if 'name="file"' not in headers:
            continue

        filename_match = re.search(r'filename="(?P<filename>[^"]+)"', headers)
        filename = filename_match.group("filename") if filename_match else ""
        return filename, file_bytes.rstrip(b"\r\n-")

    return "", b""


def main():
    parser = argparse.ArgumentParser(description="İYUK dilekçe kontrol web sunucusu")
    parser.add_argument(
        "--host",
        default=os.environ.get("HOST", HOST),
        help="Sunucu adresi",
    )
    parser.add_argument(
        "--port",
        default=int(os.environ.get("PORT", PORT)),
        type=int,
        help="Sunucu portu",
    )
    args = parser.parse_args()

    server = ThreadingHTTPServer((args.host, args.port), PetitionHandler)
    visible_host = "127.0.0.1" if args.host == "0.0.0.0" else args.host
    print(f"İYUK dilekçe kontrol aracı: http://{visible_host}:{args.port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
